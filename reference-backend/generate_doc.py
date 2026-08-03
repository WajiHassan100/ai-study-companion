import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()

# Page Margins
sections = doc.sections
for s in sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# Colors
PRIMARY = RGBColor(30, 64, 175)     # Deep Navy Blue
SECONDARY = RGBColor(71, 85, 105)   # Slate Blue
DARK_TEXT = RGBColor(30, 41, 59)    # Dark Charcoal

# Global Typography
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Arial'
font.size = Pt(10.5)
font.color.rgb = DARK_TEXT

# Title Header
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('INTERNSHIP PROGRESS & RESEARCH REPORT')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = PRIMARY
title_p.paragraph_format.space_after = Pt(4)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Personalized AI School Assistant — Multi-Agent Intelligent Education Ecosystem')
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = SECONDARY
sub_p.paragraph_format.space_after = Pt(18)

# Metadata Block
meta_table = doc.add_table(rows=2, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_cells = meta_table.rows[0].cells
meta_cells[0].text = 'Author: Software Engineering Intern'
meta_cells[1].text = 'Date: August 2026'
meta_cells2 = meta_table.rows[1].cells
meta_cells2[0].text = 'Domain: AI Autonomous Agents / EdTech'
meta_cells2[1].text = 'Framework: FastAPI + React + LangChain'

for row in meta_table.rows:
    for cell in row.cells:
        set_cell_background(cell, 'F1F5F9')
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = SECONDARY

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Helpers
def add_custom_heading(text, level=1):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = PRIMARY
    elif level == 2:
        r.font.size = Pt(11.5)
        r.font.color.rgb = SECONDARY
    return h

def add_body_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.bold = True
        r_pre.font.color.rgb = PRIMARY
    p.add_run(text)
    return p

# 1. Introduction
add_custom_heading('1. Executive Summary & Introduction', 1)
add_body_p('This report documents the research findings and engineering progress achieved during the internship on developing the Personalized AI School Assistant. The project addresses key structural challenges in modern education—such as rigid learning paces, lack of individualized feedback, and one-size-fits-all curricula—by constructing an autonomous, multi-agent AI ecosystem.')

add_custom_heading('1.1 Grounding in Literature & Research Paper Selection', 2)
add_body_p('To establish a production-ready architectural foundation, we systematically analyzed the seminal IEEE review paper "From LLM Reasoning to Autonomous AI Agents: A Comprehensive Review" (Ferrag et al., 2026). The survey provides an exhaustive review of evaluation benchmarks, framework designs (LangChain, CrewAI, Swarm, LlamaIndex), agent protocols (MCP, ACP, A2A), and multi-agent interaction topologies. This literature was selected specifically because it details the transition from passive text generation in Large Language Models (LLMs) to active, autonomous decision-making in multi-agent environments—the exact paradigm required for adaptive personalized learning.')

# 2. Key Learnings
add_custom_heading('2. Key Learnings from Research Literature', 1)

add_custom_heading('2.1 Paradigm Shift: Static LLMs vs. Autonomous AI Agents', 2)
add_body_p('As highlighted by Ferrag et al. (2026), traditional LLM applications operate in a single-turn, reactive text-generation paradigm. While effective for simple text synthesis, basic LLMs suffer from static knowledge boundaries, context window degradation, and hallucination issues.')
add_body_p('Autonomous AI Agents extend foundation models by incorporating four critical operational capabilities:', 'Core Agent Pillars: ')
add_body_p('Assigning explicit personas, domain boundaries, and output format constraints to guide reasoning.', '1. Role Specialization: ')
add_body_p('Maintaining short-term conversational context alongside persistent long-term relational and profile state.', '2. Memory & State Management: ')
add_body_p('Decomposing complex user goals into structured sub-tasks using strategies such as Chain-of-Thought (CoT) and ReAct.', '3. Planning & Multi-Step Reasoning: ')
add_body_p('Calling external APIs, executing programmatic functions, and querying backend databases dynamically.', '4. Tool Interoperability: ')

add_custom_heading('2.2 Frameworks & Communication Protocol Standards', 2)
add_body_p('The paper evaluates leading frameworks and standardized agent interaction protocols:')
add_body_p('Provide modular abstractions for prompt management, index querying, and structured tool binding.', '• LangChain & LlamaIndex: ')
add_body_p('Orchestrate teams of specialized agents with dynamic handoffs and collaborative workflows.', '• CrewAI & OpenAI Swarm: ')
add_body_p('Standardizes host-to-tool connections for secure, context-aware data delivery.', '• Model Context Protocol (MCP): ')
add_body_p('Facilitate cross-framework interoperability and structured JSON-RPC task exchanges across heterogeneous agent nodes.', '• ACP & Agent-to-Agent (A2A) Protocols: ')

# 3. Project Progress
add_custom_heading('3. Project Development Progress (Weekly Technical Report)', 1)
add_body_p('During the internship period, the project evolved from initial architectural design to a fully operational multi-agent educational platform.')

add_custom_heading('3.1 Full-Stack Architecture & Technology Selection', 2)
add_body_p('The application architecture is structured into a modern full-stack ecosystem:')
add_body_p('React 18, TypeScript, Vite, TanStack Router, Tailwind CSS, Shadcn UI components, and Lucide icons.', '• Frontend: ')
add_body_p('FastAPI (ASGI server), SQLAlchemy 2.0 (ORM), Pydantic v2 (data validation), Uvicorn, and SQLite/PostgreSQL database.', '• Backend Engine: ')
add_body_p('LangChain, OpenRouter API / Gemini API, structured JSON parsing, and custom prompt templates.', '• AI Orchestration Layer: ')

add_custom_heading('3.2 Deployed Multi-Agent Modules', 2)

# Table
table = doc.add_table(rows=5, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Agent Module', 'Status', 'Core Contribution & Functionality']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_background(hdr_cells[i], '1E40AF')
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

agent_data = [
    ('Agent #1: AI Tutor Agent (tutor_agent.py)', 'Deployed & Verified', 'Delivers 1-on-1 Socratic explanations, visual analogies, worked examples, and practice hint toggles.'),
    ('Agent #2: Assessment Profiler (profiler_agent.py)', 'Deployed & Verified', 'Grades student practice responses, computes topic mastery %, and updates StudentProfile database state.'),
    ('Agent #3: Study Planner Agent (planner_agent.py)', 'Deployed & Verified', 'Queries student weaknesses and assignment due dates to generate adaptive 7-day revision timetables.'),
    ('Agent #4: Quiz & Practice Agent (quiz_agent.py)', 'Deployed & Verified', 'Generates adaptive MCQs/flashcards, grades submissions, and auto-updates Agent #2 mastery % in DB.')
]

for row_idx, data in enumerate(agent_data, start=1):
    row_cells = table.rows[row_idx].cells
    bg = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
    for col_idx, text in enumerate(data):
        row_cells[col_idx].text = text
        set_cell_background(row_cells[col_idx], bg)
        for p in row_cells[col_idx].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# 4. Synthesis
add_custom_heading('4. Synthesis: Connecting Research to Implementation', 1)
add_body_p('The technical implementation operationalizes key concepts from the IEEE review paper:')
add_body_p('Rather than relying on a single monolithic prompt, our design separates responsibilities into specialized, single-responsibility agents, preventing prompt degradation.', '1. Role Specialization (CrewAI/Swarm): ')
add_body_p('When a student completes a quiz generated by Agent #4, the submission evaluation automatically updates the StudentProfile in Agent #2. In turn, Agent #1 (Tutor) and Agent #3 (Planner) instantly consume updated mastery scores to adjust instructional depth and schedule priorities.', '2. Closed-Loop Memory & State Updates: ')
add_body_p('Agents communicate using strictly validated Pydantic JSON schemas, reflecting tool-use best practices outlined in MCP and LangChain specifications.', '3. Structured Output & Interoperability: ')

# 5. Future Work
add_custom_heading('5. Future Work & System Roadmap', 1)
add_body_p('With the core 4-agent ecosystem fully operational, future internship phases will focus on:')
add_body_p('Indexing textbook PDFs, slides, and syllabus notes into vector stores (ChromaDB/LlamaIndex) for grounded answers with page citations.', '• Agent #5 (RAG Course Knowledge Agent): ')
add_body_p('Building class-wide analytics for educators with automated homework grading drafts and Human-in-the-Loop (HITL) confirmation controls.', '• Agent #6 (Teacher Assistant Agent): ')
add_body_p('Standardizing external tool connections and secure context delivery across agent nodes using Model Context Protocol (MCP).', '• Protocol Standardization: ')

# 6. Conclusion
add_custom_heading('6. Conclusion', 1)
add_body_p('The Personalized AI School Assistant project successfully applies state-of-the-art multi-agent research to real-world educational challenges. By integrating specialized agents for tutoring, diagnostic profiling, study planning, and automated quiz assessment, the system delivers a cohesive, stateful, and personalized learning environment. The theoretical principles derived from Ferrag et al. (2026) ensured that the resulting codebase is scalable, modular, and production-ready.')

output_path_proj = 'd:/OneDrive/OneDrive - Higher Education Commission/Desktop/Intership project/ai-study-companion/Personalized_AI_School_Assistant_Internship_Report.docx'
output_path_art = 'C:/Users/Fast Computer/.gemini/antigravity/brain/5d5fcd17-e0ed-43af-b146-ae87517a9808/Personalized_AI_School_Assistant_Internship_Report.docx'

doc.save(output_path_proj)
doc.save(output_path_art)
print('DOCX generated successfully!')
